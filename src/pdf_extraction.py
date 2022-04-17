# -*- coding: utf-8 -*-
"""
Created on Sun Apr  3 08:01:51 2022

@author: owen
"""


import pytesseract as pt
import pdf2image
import data_extraction as de
import pandas as pd

pt.pytesseract.tesseract_cmd = "C:\\Program Files/Tesseract-OCR/tesseract.exe"
"""
Code for reading pdf files with exam marks 
"""

'''
### Read in PDF with scanned-in images using pytesseract and pdf2image
'''
def pdf_image_reader(pdf_file_name):
    name = pdf_file_name.split('.pdf')[0]
    
    ### 1) Initiate to store the converted images
    pages = pdf2image.convert_from_path(pdf_path=pdf_file_name)

    ### 2) Save each page as one image into the folder 'images'        
    saved_image_path = name + "_images/"
    for i in range(len(pages)):
        pages[i].save(saved_image_path + "Page" + str(i) + '.jpg')
    

    return pages

def read_pdf():
    pass
    #out = pdf_image_reader('Sample_File_PDF_Image.pdf')
    #print(' '.join([i for i in out.split('\n') if i]))
    #C:\Users\owen\OneDrive - Eastbourne College\School analytics project\Original data files\GCSE maths marks
    #data = tabula.read_pdf(filepath, pages = 'all')
    #return data
    #print(df[0].head())

def convert_pdf_to_csv(path, filename, store = True):

    #pages = pdf_image_reader("{0}/Original data files/GCSE maths marks/{1}".format(path, filename))
    
    ### Convert pdf to image
    pdf_file = "{0}/{1}".format(path, filename)
    pages = pdf2image.convert_from_path(pdf_path=pdf_file)
    
    ### Extract the content by converting images to a list of strings 
    frames = []
    frame = pd.DataFrame()
    raw_data = []
    #DIGITS = ['0','1','2','3','4','5','6','7','8','9','0']
    for i in range(len(pages)):
        content = pt.image_to_string(pages[i])
        raw = content.split("\n\n")
        #print(raw)
        ## extract names and marks from data
        names = [r for r in raw[:20] if ':' in r]
        marks = [r for r in raw[len(names): len(names)*3] if len(r) < 5]        
        new_frame = pd.DataFrame(list(zip(names,marks)), columns = ['Name', 'Mark']) 
        frame = pd.concat([frame, new_frame], ignore_index = True)
        
        frames.append(new_frame)
        raw_data.append(raw)
    
    ### store concatenated dataframe as csv file
    if store:
        frame.to_csv("{0}/Extracted data files/GCSE_marks_csvs_from_pdfs/{1}".
                 format(de.SOURCE_DATA_DIR, filename.replace('pdf', 'csv')))
    
    return raw_data, frames
 
def overwrite_csv_from_pdf_GCSE_marks_files():    
    year = 0  # set this to one of the GCSE marks years (2016-2019) to run.  
    
    path = "{0}/Original data files/GCSE maths marks/".format(de.SOURCE_DATA_DIR)
    filename = "GCSE_scores_{0}.pdf".format(year)
    #for filename in ["GCSE_scores_2016.pdf", "GCSE_scores_2017.pdf", "GCSE_scores_2018.pdf", "GCSE_scores_2019.pdf"]:
    raw_data, dataframes = convert_pdf_to_csv(path, filename, store = False)
        
    
   
"""
** WORD DOC EXTRACTION ATTEMPT (FAILED!)
import data_extraction as de
import pandas as pd
import docx

path = "{0}/Original data files/GCSE maths marks".format(de.SOURCE_DATA_DIR)
filename = "GCSE_scores_2016.pdf"
word_filename = filename.replace("pdf", "docx")
excel_filename = filename.replace("pdf", "xlsx")



doc = docx.Document("{0}/{1}".format(path, word_filename))
#print([p.rows for p in doc.tables])

#with open("{0}/{1}".format(path, word_filename)) as file:
data = pd.read_excel("{0}/{1}".format(path, excel_filename))#,
                                 #sheet_name = year,
                           # header = [0,1]) for year in ALL_YEARS 
                            
print(data)


"""




"""
*** PDF MINER EXTRACTION ATTEMPT - FAILED SINCE SCANNED

from pdfminer.high_level import extract_text
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFParser
#text = extract_text('{0}/{1}'.format(path, filename))

#print(repr(text))#


#from pdfminer.high_level import extract_pages#
#
#for page_layout in extract_pages('{0}/{1}'.format(path, filename)):
#    print(page_layout)
#    for element in page_layout:
#        print(element)

#print("************")
#pages = extract_pages('{0}/{1}'.format(path, filename))
#print(list(pages))
#print("*****")
#for element in page1:
#    print(element)
#    for thing in element:
#        print(thing)
        





from io import StringIO
output_string = StringIO()

with open('{0}/{1}'.format(path, filename), 'rb') as in_file:
    parser = PDFParser(in_file)
    doc = PDFDocument(parser)
    rsrcmgr = PDFResourceManager()
    device = TextConverter(rsrcmgr, output_string, laparams=LAParams(line_overlap = 10, char_margin = 1000, 
                                                                     boxes_flow = -1.0, all_texts = True))
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    for page in PDFPage.create_pages(doc):
        print(page)
        interpreter.process_page(page)

output_string.seek(0)
#print(output_string.read())
#print(output_string.getvalue())


"""